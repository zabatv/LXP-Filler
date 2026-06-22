import io
import os
import json
from datetime import datetime, timezone
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
from urllib.parse import urlparse

import requests
from docx import Document
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import HTMLResponse, Response, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel
from starlette.requests import Request

app = FastAPI(title="LXP Journal Filler Web")

API_URL = "https://api.newlxp.ru/graphql"

templates = Jinja2Templates(directory="templates")


# ---------- GraphQL helper ----------

def graphql(token: str, query: str, variables: dict = None) -> dict:
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    body = {"query": query}
    if variables:
        body["variables"] = variables
    resp = requests.post(API_URL, headers=headers, json=body, timeout=30)
    data = resp.json()
    if data.get("errors"):
        raise HTTPException(status_code=400, detail=data["errors"][0]["message"])
    return data["data"]


# ---------- Pages ----------

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse("login.html", {"request": request})


@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard(request: Request):
    return templates.TemplateResponse("dashboard.html", {"request": request})


# ---------- Auth API ----------

class LoginInput(BaseModel):
    email: str
    password: str


@app.post("/api/auth/login")
async def login(input: LoginInput):
    query = """
        query SignIn($input: SignInInput!) {
            signIn(input: $input) { accessToken }
        }
    """
    resp = requests.post(API_URL, json={
        "query": query,
        "variables": {"input": {"email": input.email.strip(), "password": input.password}}
    }, timeout=30)
    data = resp.json()
    errors = data.get("errors")
    if errors:
        raise HTTPException(status_code=401, detail=errors[0].get("message", "Ошибка входа"))
    token = data.get("data", {}).get("signIn", {}).get("accessToken")
    if not token:
        raise HTTPException(status_code=401, detail="Токен не получен")
    return {"token": token}


@app.post("/api/auth/check")
async def check_token(request: Request):
    body = await request.json()
    token = body.get("token", "")
    if not token:
        raise HTTPException(status_code=400, detail="Токен не предоставлен")
    if token.startswith("Bearer "):
        token = token[7:]
    try:
        graphql(token, "query { getMe { id } }")
        return {"valid": True}
    except HTTPException:
        raise HTTPException(status_code=401, detail="Токен недействителен")


# ---------- Data API ----------

@app.get("/api/suborganizations")
async def get_suborganizations(token: str = ""):
    if not token:
        raise HTTPException(status_code=401, detail="Требуется токен")
    data = graphql(token, """
        query {
            getMe {
                assignedSuborganizations {
                    suborganizationId
                    suborganization { id name organizationId }
                }
            }
        }
    """)
    items = data["getMe"]["assignedSuborganizations"]
    seen = {}
    result = []
    for item in items:
        key = item["suborganizationId"]
        if key not in seen:
            seen[key] = True
            result.append(item)
    return {"items": result}


@app.get("/api/groups")
async def get_groups(token: str = "", org_id: str = "", suborg_id: str = "", study_period_id: str = ""):
    if not token:
        raise HTTPException(status_code=401, detail="Требуется токен")
    if study_period_id:
        data = graphql(token, f"""
            query {{
                learningGroupsByStudyPeriodIdAndSuborganizationId(input: {{
                    studyPeriodId: "{study_period_id}"
                    suborganizationId: "{suborg_id}"
                }}) {{ id name }}
            }}
        """)
        return {"items": data["learningGroupsByStudyPeriodIdAndSuborganizationId"]}
    data = graphql(token, f"""
        query {{
            getLearningGroups(input: {{ organizationId: "{org_id}" suborganizationId: "{suborg_id}" isArchived: false }}) {{
                id name
            }}
        }}
    """)
    return {"items": data["getLearningGroups"]}


@app.get("/api/study-periods")
async def get_study_periods(token: str = "", org_id: str = ""):
    if not token:
        raise HTTPException(status_code=401, detail="Требуется токен")
    data = graphql(token, f"""
        query {{
            studyPeriods(input: {{ filters: {{ organizationId: "{org_id}" }} }}) {{
                id name startDate endDate
            }}
        }}
    """)
    now = datetime.now(timezone.utc)
    items = sorted(data["studyPeriods"], key=lambda x: x["startDate"])
    # Mark current period (today falls within startDate-endDate)
    for sp in items:
        try:
            end = datetime.fromisoformat(sp["endDate"].replace("Z", "+00:00"))
            start = datetime.fromisoformat(sp["startDate"].replace("Z", "+00:00"))
            sp["isCurrent"] = start <= now <= end
        except Exception:
            sp["isCurrent"] = False
    return {"items": items}


@app.get("/api/disciplines")
async def get_disciplines(token: str = "", group_id: str = "", semester: str = "1"):
    if not token:
        raise HTTPException(status_code=401, detail="Требуется токен")
    data = graphql(token, f"""
        query {{
            disciplinesByGroups(input: {{ groupIds: ["{group_id}"] }}) {{
                id name code
                teachers {{ user {{ lastName firstName middleName }} }}
            }}
        }}
    """)
    return {"items": data["disciplinesByGroups"]}


@app.get("/api/students")
async def get_students(token: str = "", group_id: str = "", disc_id: str = "", study_period_id: str = ""):
    if not token:
        raise HTTPException(status_code=401, detail="Требуется токен")
    data = graphql(token, f"""
        query {{
            searchStudentsInLearningGroup(input: {{
                filters: {{ learningGroupId: "{group_id}", isExpelled: false }}
            }}) {{
                items {{ id user {{ lastName firstName middleName }} }}
            }}
        }}
    """)
    students = data["searchStudentsInLearningGroup"]["items"]

    # Find teacher
    teacher_name = ""
    try:
        disc_data = graphql(token, f"""
            query {{
                disciplinesByGroups(input: {{ groupIds: ["{group_id}"] }}) {{
                    id teachers {{ user {{ lastName firstName middleName }} }}
                }}
            }}
        """)
        for d in disc_data["disciplinesByGroups"]:
            if d["id"] == disc_id and d["teachers"]:
                t = d["teachers"][0]["user"]
                teacher_name = f"{t['lastName']} {t['firstName']} {t.get('middleName', '')}".strip()
                break
    except Exception:
        pass

    # Pre-build name map (names already known from searchStudentsInLearningGroup)
    name_map = {s["id"]: f"{s['user']['lastName']} {s['user']['firstName']} {s['user'].get('middleName', '')}".strip() for s in students}

    # Fetch grades concurrently
    def get_grade(student_id, idx):
        try:
            name = name_map.get(student_id, "Ошибка")
            grade = ""
            has_retake = False
            retake_grade = ""
            retake_score = ""
            if study_period_id:
                sd_data = graphql(token, f"""
                    query {{
                        searchStudentDisciplines(input: {{
                            studentId: "{student_id}"
                            filters: {{ studyPeriodId: "{study_period_id}" }}
                        }}) {{
                            disciplineId disciplineGrade hasRetake retakeDisciplineGrade retakeScore
                        }}
                    }}
                """)
                for sd in sd_data["searchStudentDisciplines"]:
                    if sd["disciplineId"] == disc_id:
                        grade = sd["disciplineGrade"] or ""
                        has_retake = sd.get("hasRetake", False)
                        retake_grade = sd.get("retakeDisciplineGrade", "")
                        retake_score = sd.get("retakeScore", "")
                        break
            else:
                gdata = graphql(token, f"""
                    query {{
                        getUserById(input: {{ userId: "{student_id}" }}) {{
                            student {{ studentDiscipline(disciplineId: "{disc_id}") {{ disciplineGrade }} }}
                        }}
                    }}
                """)
                sd = gdata["getUserById"]["student"]["studentDiscipline"]
                grade = sd["disciplineGrade"] if sd else ""
            GRADE_MAP = {"TWO": "2", "THREE": "3", "FOUR": "4", "FIVE": "5"}
            if has_retake and retake_grade in GRADE_MAP:
                grade = GRADE_MAP[retake_grade]
            return {"id": student_id, "name": name, "grade": grade, "hasRetake": has_retake, "retakeGrade": retake_grade, "retakeScore": retake_score, "idx": idx}
        except Exception:
            return {"id": student_id, "name": name_map.get(student_id, "Ошибка"), "grade": "", "idx": idx}

    results = []
    with ThreadPoolExecutor(max_workers=10) as executor:
        futures = {executor.submit(get_grade, s["id"], i): s for i, s in enumerate(students)}
        for future in as_completed(futures):
            try:
                results.append(future.result(timeout=10))
            except Exception:
                pass
    results.sort(key=lambda x: x["idx"])

    return {"items": results, "teacher_name": teacher_name, "count": len(results)}


# ---------- DOCX Fill API ----------

@app.post("/api/docx/fill")
async def fill_docx(
    file: UploadFile = File(...),
    token: str = Form(...),
    group_name: str = Form(""),
    disc_name: str = Form(""),
    teacher_name: str = Form(""),
    students_json: str = Form("[]"),
):
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Файл должен быть в формате DOCX")

    content = await file.read()
    doc = Document(io.BytesIO(content))
    students = json.loads(students_json)

    student_idx = 0
    for table in doc.tables:
        for row in table.rows:
            has_marker = any("%n" in cell.text or "%q" in cell.text for cell in row.cells)
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if "%g" in run.text:
                            run.text = run.text.replace("%g", group_name)
                        if "%d" in run.text:
                            run.text = run.text.replace("%d", disc_name)
                        if "%t" in run.text:
                            run.text = run.text.replace("%t", teacher_name)
                        if has_marker and student_idx < len(students):
                            if "%n" in run.text:
                                run.text = run.text.replace("%n", students[student_idx]["name"])
                            if "%q" in run.text:
                                grade = str(students[student_idx]["grade"]) if students[student_idx]["grade"] else "N/A"
                                run.text = run.text.replace("%q", grade)
            if has_marker:
                student_idx += 1

    for para in doc.paragraphs:
        for run in para.runs:
            if "%g" in run.text:
                run.text = run.text.replace("%g", group_name)
            if "%d" in run.text:
                run.text = run.text.replace("%d", disc_name)
            if "%t" in run.text:
                run.text = run.text.replace("%t", teacher_name)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=filled_{file.filename}"},
    )


# ---------- Example DOCX download ----------

@app.get("/example")
async def download_example():
    file_path = os.path.join(os.path.dirname(__file__), "static", "example.docx")
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="example.docx not found")
    with open(file_path, "rb") as f:
        content = f.read()
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=example.docx"},
    )


if __name__ == "__main__":
    import sys
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    print(f"Starting LXP Journal Filler on port {port}", flush=True)
    sys.stdout.flush()
    uvicorn.run(app, host="0.0.0.0", port=port, log_level="info")
